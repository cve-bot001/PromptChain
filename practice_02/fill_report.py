#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(
    "D:\\文档\\班级\\作业\\人工智能提示工程\\AI_Program\\practice_02\\学号-班级-姓名-第x次实验报告.docx"
)

for para in doc.paragraphs:
    if "[学 号]" in para.text:
        para.text = "学    号： 22403010125"
    if "[姓 名]" in para.text:
        para.text = "姓    名： 张三"
    if "[班 级]" in para.text:
        para.text = "班    级： 24401"

for i, para in enumerate(doc.paragraphs):
    if "实验目的" in para.text and i < 10:
        para.text = "【实验目的】\n掌握Function Call工具调用功能，理解LLM如何调用外部工具实现文件操作。"

for i, para in enumerate(doc.paragraphs):
    if "实验内容" in para.text and i < 10:
        para.text = "【实验内容】\n1. 开发基于Function Call的5个文件操作工具\n2. 实现终端聊天客户端，支持流式输出和历史上下文\n3. 让LLM理解自然语言指令并调用相应工具"

for i, para in enumerate(doc.paragraphs):
    if "实验过程" in para.text and "1." in para.text:
        para.text = "【实验过程】\n1. 创建tool_agent.py，实现5个Function Call工具：\n   - list_files_with_info：列出目录下文件及属性\n   - rename_file：重命名文件\n   - delete_file：删除文件\n   - create_file：创建新文件并写入内容（支持覆盖）\n   - read_file_content：读取文件内容\n2. 创建chat_terminal.py终端聊天客户端：\n   - 支持流式输出(streaming)\n   - 历史聊天记录自动添加到上下文\n   - Ctrl+C退出\n3. 配置.env文件，设置BASE_URL和MODEL\n4. 启动本地llama.cpp server加载模型\n5. 运行python practice_02/tool_agent.py测试"

for i, para in enumerate(doc.paragraphs):
    if "实验结果" in para.text and "1." in para.text:
        para.text = '【实验结果】\n成功实现5个Function Call工具调用：\n- list_files_with_info：列出目录文件列表及大小、时间属性\n- rename_file：成功重命名文件（如newfile.txt->renamed.txt）\n- delete_file：成功删除指定文件\n- create_file：成功创建文件并写入内容，文件已存在时自动覆盖\n- read_file_content：成功读取文件内容并返回\n\n测试指令：\n- "list files in practice_02" -> 列出文件\n- "view test.txt" -> 读取文件\n- "delete test.txt" -> 删除文件\n- "create newfile.txt with content hello" -> 创建文件\n- "rename old.txt to new.txt" -> 重命名文件'

doc.save(
    "D:\\文档\\班级\\作业\\人工智能提示工程\\AI_Program\\practice_02\\学号-班级-姓名-第x次实验报告_已填写.docx"
)
print("Done! Saved as 学号-班级-姓名-第x次实验报告_已填写.docx")
